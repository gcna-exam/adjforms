'''
Parse returned (docx format) adjudication forms 
Flag missing grades
Create json dictionary of results
Create pdf summary of overall pass/fail marks, for committee use
Create csv summary of individual piece and overall marks, for board use

Created by M. Pan, fall 2020
Edited spring 2021, fall 2021, spring 2022


USAGE:

Collect all .docx files returned by jurors, and a copy of parse_adjforms.py, into one directory.
Make that directory the current working directory.

$ python3 parse_adjforms.py

*** NB : move all extraneous juror files *out* of directory before running
*** NB : close all docx files to be parsed before running


THINGS TO BE EDITED EACH YEAR:

-- "Parameters"
'''

import copy
import glob
import json
import os
import random
import subprocess

from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
#from docx2pdf import convert


############################################
##### Parameters --- !!!CHANGE THIS!!! #####
############################################

# year in which exam cycle finishes
examyear = '2022'

# list of all jurors in committee
jurors = ['hunsberger','lee','lehrer','lens','lukyanova','macoska','tam']

# list of all candidates (numbers) with recordings submitted
candidates = ['1','3','4','5','6','7','8','9','10','11','12','13']

# list of *voting* jurors only
voting = ['hunsberger','lehrer','lens','macoska','tam']

# dictionary of juror recusals, item format <candidate number> : [list of recused jurors]
conflict = {'13':['hunsberger']}#{'3':['lens']}

# number of adjudications needed for a full slate of votes
numrequired = 5

# 'prelim' for pre-juror discussing, 'final' for post-juror discussion
labelstr = 'final'


############################
##### Helper functions #####
############################

def insensitive_glob(pattern):

    # return all filenames containing case-insensitive version of pattern
    
    def either(c):
        return '[%s%s]' % (c.lower(), c.upper()) if c.isalpha() else c
    
    return glob.glob(''.join(map(either, pattern)))


def get_candnumber(adjform):

    # extract candidate number from adjform Document

    return adjform.paragraphs[3].text.split('\t')[0].split('.')[-1].strip()


def get_pieces(adjform):

    # extract piece names from adjform Document

    piecelines = [ind-1 for ind,par in enumerate(adjform.paragraphs) if par.text == 'piece\t\t\t\t\t\t\tcomposer']

    return [adjform.paragraphs[ind].text.split('\t')[0] for ind in piecelines]


def get_grades(adjform):

    # extract grades from adjform Document for individual pieces
    # OLD, DEPRECATED IN FAVOR OF get_pf_grades

    gradelines = [par.text for par in adjform.paragraphs if par.text[:9] == 'Candidate']

    grades = []
    for line in gradelines:
        gradestr = line.split('\t')[-1].split('Over')[0].strip('_')

        try:
            thisgrade = float(gradestr)
        except:
            # check for fractions with "/" divider
            if '/' in gradestr:
                if ' ' in gradestr:
                    wholegrade = gradestr.split()[0]
                    if wholegrade.isnumeric():
                        fracgrade = gradestr.split()[1].strip().split('/')
                        if len(fracgrade) == 2 and fracgrade[0].isnumeric() and fracgrade[1].isnumeric():
                            thisgrade = int(wholegrade) + float(fracgrade[0])/float(fracgrade[1])
                        else:
                            thisgrade = float(wholegrade)
                    else:
                        thisgrade = 0
                else:
                    gradeparts = gradestr.split('/')
                    wholegrade = gradeparts[0][:-1]
                    fracgrade = [gradeparts[0][-1],gradeparts[1]]
                    if wholegrade.isnumeric():
                        if fracgrade[0].isnumeric() and fracgrade[1].isnumeric():
                            thisgrade = int(wholegrade) + float(fracgrade[0])/float(fracgrade[1])
                        else:
                            thisgrade = float(wholegrade)
                    else:
                        thisgrade = 0
            else:
                # check for "+","-" in grade
                trygrade = gradestr.replace('+','').replace('-','')
                if trygrade.isnumeric():
                    if '+' in gradestr:
                        offset = 0.3
                    elif '-' in gradestr:
                        offset = -0.3
                    thisgrade = float(trygrade)+offset
                else:
                    thisgrade = 0
                
        grades.append(thisgrade)

    return grades
                

def get_pf_grades(adjform):

    # extract pass/fail grades from adjform Document for individual pieces

    gradelines = [par for par in adjform.paragraphs if par.text[:9] == 'Candidate']

    grades = []
    for line in gradelines:
        gradestr = line.text.split('Rating:')[1]
        if 'not' not in gradestr.lower():
            passmark = 'x'
            failmark = ''
        elif gradestr.lower().count('pass') == 1:
            passmark = ''
            failmark = 'x'
        else:
            passmark = gradestr.split('passing')[0].strip().replace('_','')
            failmark = line.text.split('passing')[1].split('not')[0].strip().replace('_','')

        if passmark and not failmark:
            thisgrade = 1
        elif failmark and not passmark:
            thisgrade = -1
        elif passmark and failmark:
            if passmark.lower() in ['y','yes'] and failmark.lower in ['n','no']:
                thisgrade = 1
            elif passmark.lower() in ['n','no'] and failmark.lower in ['y','yes']:
                thisgrade = -1
            else:
                thisgrade = check_grade_formatting(line)
        else:
            thisgrade = check_grade_formatting(line)
            
        grades.append(thisgrade)

    return grades


def check_grade_formatting(par):

    # grade info is not contained in text, check formatting to see if can decipher

    passingrun = [run for run in par.runs if 'pass' in run.text.lower() and 'not' not in run.text.lower()]
    notpassingrun = [run for run in par.runs if 'pass' in run.text.lower() and 'not' in run.text.lower()]
    if len(passingrun) == 1 and len(notpassingrun) == 1:
        if passingrun[0].font.highlight_color and not notpassingrun[0].font.highlight_color:
            grade = 1
        elif not passingrun[0].font.highlight_color and notpassingrun[0].font.highlight_color:
            grade = -1
        elif passingrun[0].font.underline and not notpassingrun[0].font.underline:
            grade = 1
        elif not passingrun[0].font.underline and notpassingrun[0].font.underline:
            grade = -1
        elif not passingrun[0].font.strike and notpassingrun[0].font.strike:
            grade = 1
        elif passingrun[0].font.strike and not notpassingrun[0].font.strike:
            grade = -1
        else:
            grade = 0
    else:
        grade = 0

    return grade

            

def get_req(adjform):

    # extract req/non-req piece order from adjform Document
    # 1 = required , 0 = non-required

    reqlines = [par.text for par in adjform.paragraphs if par.text[:8] == 'Required']

    reqlist = []
    for line in reqlines:
        reqstr = line.split('\t')[1].strip()
        if reqstr == 'Yes':
            reqlist.append(1)
        elif reqstr == 'No':
            reqlist.append(0)
        else:
            reqlist.append(-1)

    return reqlist
        

def get_overall(adjform):

    # extract overall pass/fail from adjform Document

    overall = ''
    
    try:
        overallline = adjform.tables[0].rows[1].cells[1].paragraphs[1]
    except:
        try:
            overallline = adjform.tables[0].rows[1].cells[1].paragraphs[0]
        except:
            print("can't get line with overall mark")

    # if there are no formatting details in this line, just check text
    if len(overallline.runs) == 1:
        if 'not pass' in overallline.text and overallline.text.count('pass') == 1:
            overall = 'fail'
        elif 'pass' in overallline.text and 'not' not in overallline.text:
            overall = 'pass'
        else:
            overall = ""
            
    # but if there are formatting details the meaning might be different from that of the text string
    elif len(overallline.runs) >= 2:
        passrun = [x for x in overallline.runs if 'pass' in x.text and 'not' not in x.text]
        notpassrun = [x for x in overallline.runs if 'not' in x.text]
        if notpassrun and passrun:
            if notpassrun[0].font.bold and not passrun[0].font.bold:
                overall = 'fail'
            elif notpassrun[0].font.underline and not passrun[0].font.underline:
                overall = 'fail'
            elif notpassrun[0].font.strike:
                overall = 'pass'
            elif notpassrun[0].font.highlight_color and not passrun[0].font.highlight_color:
                overall = 'fail'
            else:
                overall = 'pass'
        elif notpassrun:
            overall = 'fail'
        else:
            overall = 'pass'
            
    else:
        overall = "can't parse"

    return overall
    

def record_grades(thisdict,pieces,grades):

    # record grades for this juror in results dictionary
    
    for ind in range(len(pieces)):

        piece = pieces[ind]
        grade = grades[ind]
        
        if grade == 0:
            print('cand '+candidate+'/'+juror+'/'+piece+' : missing grade')
            gradecode = 0
        elif grade == 1:
            gradecode = 'pass'
        elif grade == -1:
            gradecode = 'fail'
        else:
            print('cand '+candidate+'/'+juror+'/'+piece+' : unknown grade '+str(grade))
            gradecode = str(grade)
            
        if piece not in thisdict.keys():
            thisdict[piece] = {juror:gradecode}
        else:
            thisdict[piece][juror] = gradecode

    return thisdict


def record_overall(thisdict,overall,req_outcomes):

    # record overall grade, check for consistency

    if overall == 'pass':
        thisdict['pass'].append(juror)
        #if any(prod>0 and prod<3 for prod in req_outcomes):
        if any(prod < 0 for prod in req_outcomes):
            print('cand '+candidate+'/'+juror+' : overall pass, failed req')
    elif overall == 'fail':
        thisdict['fail'].append(juror)
        #if all(prod==0 or prod>=3 for prod in req_outcomes):
        if all(prod >= 0 for prod in req_outcomes):
            print('cand '+candidate+'/'+juror+' : overall fail, all req passed')
    else:
        print('cand '+candidate+'/'+juror+' : overall mark "'+overall+'"')

    return thisdict


def record_repvotes(repform,results,juror):

    # read juror's votes from repertoire form, record in results dictionary
    
    for rownum in range(len(repform.tables[0].rows) - 1):
        thiscand,piece1,piece2,piece3 = [repform.tables[0].row_cells(rownum+1)[ind].text for ind in range(4)]
        
        if thiscand and any([piece1,piece2,piece3]):
            thisdict = copy.deepcopy(results[thiscand]['repertoire'])
            results[thiscand]['repertoire'] = write_repvotes(thisdict,thiscand,juror,piece1,piece2,piece3)

    return results


def write_repvotes(thisdict,thiscand,juror,piece1,piece2,piece3):

    # record a single juror's repertoire piece votes for a single candidate

    pieces = [piece1,piece2,piece3]
    for ind in [1,2,3]:
        piece = pieces[ind-1]
        if piece:
            if piece not in thisdict.keys():
                thisdict[piece] = {ind:[juror]}
            elif ind not in thisdict[piece].keys():
                thisdict[piece][ind] = [juror]
            else:
                thisdict[piece][ind].append(juror)
        else:
            print(thiscand+' : missing repertoire piece '+str(ind))

    return thisdict


def record_reqvotes(reqform,results,juror):

    # read juror's votes from required piece form, record in results dictionary
    
    parind = [ind for ind,par in enumerate(reqform.paragraphs) if 'Candidate' in par.text]
    
    for ind in parind:
        
        thiscand,techpiece,techmark = reqform.paragraphs[ind].text.split('\t')
        _,exppiece,expmark = reqform.paragraphs[ind+1].text.split('\t')
        
        thiscand = thiscand.split()[-1]
        techmark = techmark.replace('_','')
        expmark = expmark.replace('_','')
        techruns = reqform.paragraphs[ind].runs
        expruns = reqform.paragraphs[ind+1].runs
        
        thisdict = copy.deepcopy(results[thiscand]['required'])
        results[thiscand]['required'] = write_reqvote(thiscand,thisdict,juror,techpiece,exppiece,techmark,expmark,techruns,expruns)
            
    return results


def write_reqvote(cand,thisdict,juror,techpiece,exppiece,techmark,expmark,techruns,expruns):

    # record juror's required piece vote for a single candidate

    if techpiece not in thisdict.keys():
        thisdict[techpiece] = []
    if exppiece not in thisdict.keys():
        thisdict[exppiece] = []
    
    if techmark and not expmark:
        thisdict[techpiece].append(juror)
    elif expmark and not techmark:
        thisdict[exppiece].append(juror)
    elif techmark and expmark:
        if '1' in techmark and '1' not in expmark:
            thisdict[techpiece].append(juror)
        elif '1' in expmark and '1' not in techmark:
            thisdict[exppiece].append(juror)
        else:
            print(cand+" : can't parse required piece vote")
    else:
        vote = check_vote_formatting(techpiece,exppiece,techruns,expruns)
        if vote == 'tech':
            thisdict[techpiece].append(juror)
        elif vote == 'exp':
            thisdict[exppiece].append(juror)
        else:
            print(cand+" : can't parse required piece vote")

    return thisdict


def check_vote_formatting(techpiece,exppiece,techruns,expruns):

    # if can't understand required piece vote using text, check font/formatting

    techtitle = [run for run in techruns if any(word in run.text for word in techpiece.split())]
    exptitle = [run for run in expruns if any(word in run.text for word in exppiece.split())]
    
    if techtitle and exptitle:
        if any(run.font.strike for run in techtitle) and not any(run.font.strike for run in exptitle):
            result = 'exp'
        elif any(run.font.strike for run in exptitle) and not any(run.font.strike for run in techtitle):
            result = 'tech'
        elif any(run.font.underline for run in techtitle) and not any(run.font.underline for run in exptitle):
            result = 'tech'
        elif any(run.font.underline for run in exptitle) and not any(run.font.underline for run in techtitle):
            result = 'exp'
        elif any(run.font.bold for run in techtitle) and not any(run.font.bold for run in exptitle):
            result = 'tech'
        elif any(run.font.bold for run in exptitle) and not any(run.font.bold for run in techtitle):
            result = 'exp'
        elif any(run.font.highlight_color for run in techtitle) and not any(run.font.highlight_color for run in exptitle):
            result = 'tech'
        elif any(run.font.highlight_color for run in exptitle) and not any(run.font.highlight_color for run in techtitle):
            result = 'exp'
        else:
            result = ''
    else:
        result = ''

    return result

            
def make_jurorsummary(results,jurors,voting,conflicts):

    # construct summary of overall grades for committee reference

    # find nonvoting jurors, set up random choice of one for each candidate
    altjurors = [x for x in jurors if x not in voting]
    if altjurors:
        altchoices = random.Random(int(examyear)).choices(range(len(altjurors)),[1]*len(altjurors),k=len(candidates))
    else:
        altchoices = []

    votingsummary = {}
    
    for candidate in candidates:

        # check which jurors marked this candidate
        thisjurors = results[candidate]['pass'] + results[candidate]['fail']
        thisaltjurors = [x for x in altjurors if x in thisjurors]

        # check for juror recusals
        if candidate in conflict.keys():
            recuse = conflict[candidate]
        else:
            recuse = []

        # determine voting jurors for this candidate by picking altjuror if any recusals
        if len(recuse) > len(thisjurors)-numrequired:
            
            # there are a lot of recusals or AWOL voting jurors, just use all jurors possible
            print('cand '+candidate+' : not enough jurors')
            thisvoting = [x for x in thisjurors if x not in recuse]
            
        elif set(voting).intersection(set(recuse)):
            
            # there are recusals and there are enough alternates
            if len(thisaltjurors) == 1:
                thisvoting = [x for x in voting if x not in recuse and x in thisjurors] + thisaltjurors
                if len(thisvoting) < numrequired:
                    print('cand '+candidate+' : not enough jurors')

            elif len(thisaltjurors) > 1:
                thisvoting = [x for x in voting if x not in recuse and x in thisjurors] + [thisaltjurors[altchoices[0]]]
                if len(thisvoting) < numrequired:
                    thisvoting.append(altjurors[1+altchoices[0] % 2])
                    
            else:
                thisvoting = [x for x in voting if x not in recuse and x in thisjurors]

            if altchoices:
                altchoices = altchoices[1:]

        else:
            
            # in effect there are no recusals
            thisvoting = [juror for juror in voting if juror in thisjurors]

            # but if there are not enough voting jurors, still have to use alternates
            if numrequired != len(thisvoting):
                okaltjurors = [x for x in thisaltjurors if x not in recuse]
                if numrequired-len(thisvoting) >= len(okaltjurors):
                    thisvoting += okaltjurors
                else:
                    thisvoting += okaltjurors[altchoices[0]]

            if len(thisvoting) < numrequired:
                print('cand '+candidate+' : not enough jurors')

            if altchoices:
                altchoices = altchoices[1:]
            

        # store vote tallies
        numpass = len([x for x in results[candidate]['pass'] if x in thisvoting])
        numfail = len([x for x in results[candidate]['fail'] if x in thisvoting])
        votingsummary[candidate] = {'pass/fail':[numpass,numfail],'voting':thisvoting}

        # count required piece votes
        # select piece with most votes, unless there is a tie
        reqcountvotes = [[piece,len(results[candidate]['required'][piece])] for piece in results[candidate]['required'].keys()]
        
        if reqcountvotes:
            reqselect = max(reqcountvotes,key=lambda pair:pair[1])
            testreqselect = [pair[0] for pair in reqcountvotes if pair[1] == reqselect[1]]
        else:
            reqselect = []
            testreqselect = []
            
        if len(testreqselect) > 1:
            print('candidate '+candidate+', '+', '.join(testreqselect)+' : required piece tie vote')
            reqpiece = ', '.join(testreqselect)+' (tie)'
        else:
            try:
                reqpiece = reqselect[0]
            except:
                reqpiece = ''

        votingsummary[candidate]['reqpiece'] = reqpiece

        # count repertoire piece votes
        # weighting 3 points for 1st choice, 2 for 2nd choice, 1 for 3rd choice
        # select piece with most points, unless there is a tie
        thisrepdict = results[candidate]['repertoire']
        for piece in thisrepdict.keys():
            if 1 not in thisrepdict[piece].keys():
                thisrepdict[piece][1] = []
            if 2 not in thisrepdict[piece].keys():
                thisrepdict[piece][2] = []
            if 3 not in thisrepdict[piece].keys():
                thisrepdict[piece][3] = []
                
        repcountvotes = [[piece,len(thisrepdict[piece][1])*3
                          +len(thisrepdict[piece][2])*2
                          +len(thisrepdict[piece][3])] for piece in thisrepdict.keys()]

        if repcountvotes:
            repselect = max(repcountvotes,key=lambda pair:pair[1])
            testrepselect = [pair[0] for pair in repcountvotes if pair[1] == repselect[1]]
        else:
            repselect = []
            testrepselect = []
        
        if len(testrepselect) > 1:
            print('candidate '+candidate+' -- '+', '.join(testrepselect)+' : repertoire piece tie vote')
            reppiece = ', '.join(testrepselect)+' (tie)'
        else:
            try:
                reppiece = repselect[0]
            except:
                reppiece = ''

        votingsummary[candidate]['reppiece'] = reppiece

        if abs(numpass-numfail) <= 1 or 'prelim' not in labelstr:
            make_candidate_pdf(candidate,thisjurors)

            
    with open(examyear+'votingsummary.json','w') as fh:
        json.dump(votingsummary,fh,indent=4,sort_keys=True)


    # create latex document with grade summary
    preamble = [
        '\\documentclass[10pt]{article}\n', \
        '\n', \
        '\\usepackage{parskip,array}\n', \
        '\\usepackage[scaled=.9]{helvet}\n', \
        '\\usepackage[T1]{fontenc}\n', \
        '\n', \
        '\\addtolength{\\topmargin}{-.9in}\n', \
        '\\addtolength{\\oddsidemargin}{0in}\n', \
        '\\addtolength{\\oddsidemargin}{-1in}\n', \
        '\\addtolength{\\textwidth}{2in}\n', \
        '\\addtolength{\\textheight}{1.7in}\n', \
        '\n', \
        '\\renewcommand\\familydefault{\\sfdefault}\n', \
        '\n', \
        '\\begin{document}\n', \
        '\n', \
        '\\begin{LARGE}\n', \
        '\\noindent {\\bf '+examyear+' '+labelstr+' recording stage results}\\bigskip\\hfill\n', \
        '\\end{LARGE}\n', \
        '\n', \
        '\\begin{large}\n', \
        ]

    votetable = ['\\begin{tabular}{cc@{\\hspace{1.8in}}cc}\n', \
                 '{\\bf Candidate}& {\\bf pass/fail}\\makebox[0in][l]{ (voting jurors only)}& {\\bf required}& {\\bf repertoire}\\smallskip\\\\\n']

    for candidate in candidates:

        numpass = votingsummary[candidate]['pass/fail'][0]
        numfail = votingsummary[candidate]['pass/fail'][1]
        req = votingsummary[candidate]['reqpiece']
        rep = votingsummary[candidate]['reppiece']
        
        votetable += [candidate+'& '+str(numpass)+'/'+str(numfail)+'& '+req+'& '+rep+'\\\\\n']
        
    votetable += ['\\end{tabular}\n\n']

    votelines = preamble + votetable + ['\\end{large}\n','\\end{document}\n']
    
    votefile = examyear+'_'+labelstr+'_recording_summary.tex'
    with open(votefile,'w') as fh:
        for line in votelines:
            _ = fh.write(line)

    subprocess.run('pdflatex '+votefile+' > pdflatex.out',shell=True)


def make_candidate_pdf(candidate,thisjurors):

    # create pdf of all adjudications for juror review

    composer = ''
    for juror in thisjurors:
        
        filelist = insensitive_glob('*'+juror+'*docx')
        thisformname = [form for form in filelist if 'candidate'+candidate+'_' in form]

        if not thisformname:
            print('cand '+candidate+'/'+juror+' : strange name for form?')
            thisformname = ''
            
        else:
            thisform = Document(thisformname[0])
            numpar = len(thisform.paragraphs)
            for ind in range(numpar):
                par = thisform.paragraphs[ind]
                if 'Juror' in par.text and 'Signature' in par.text:
                    run = par.add_run()
                    run.add_break(WD_BREAK.PAGE)
                if ind == numpar-1:
                    if len(par.runs) != 0:
                        run = par.add_run()
                        run.add_break(WD_BREAK.PAGE)
            if not composer:
                composer = Composer(thisform)
            else:
                try:
                    composer.append(thisform)
                except:
                    print(thisformname[0]+' error in make_candidate_pdf composer.append')

    composer.save(examyear+'_candidate'+candidate+'_all.docx')
    #convert(examyear+'_candidate'+candidate+'_all.docx')
    

def make_boardsummary(results):

    # create csv for board convenience

    csvlist = ['candidate,piece,'+','.join(jurors)+',,overall,'+','.join(jurors)]
    #csvlist = ['candidate,piece,'+','.join(jurors)+',range,avg,,overall,'+','.join(jurors)]

    for candidate in candidates:

        overallstr = ''

        pieces = [key for key in results[candidate].keys() if key not in ['pass','fail','required','repertoire']]
        for piece in pieces:
            
            thisline = candidate + ',' + piece.replace(',','') + ','
            
            for juror in jurors:
                if juror in results[candidate][piece].keys():
                    thisline += str(results[candidate][piece][juror]) + ','
                else:
                    thisline += '0,'
                    
            #thispiecegrades = [v for k,v in results[candidate][piece].items() if v != 0]
            #if thispiecegrades:
            #    graderange = max(thispiecegrades) - min(thispiecegrades)
            #    gradeavg = sum(thispiecegrades)/len(thispiecegrades)
            #else:
            #    graderange = 0
            #    gradeavg = 0
            #    
            #thisline += str(graderange) + ',' + '{:4.2f}'.format(gradeavg) + ','*(len(jurors)+2)
            thisline += ','*(len(jurors)+1)
            csvlist.append(thisline)

        for juror in jurors:
            if juror in results[candidate]['pass']:
                overallstr += ',P'
            elif juror in results[candidate]['fail']:
                overallstr += ',F'
            else:
                overallstr += ','

        csvlist[-1] = csvlist[-1].replace(','*(len(jurors)+2),',,'+overallstr)    
        
    with open(examyear+'_'+labelstr+'_exam_grade_summary.csv','w') as fh:
        for line in csvlist:
            _ = fh.write(line+'\n')
        

        
##########################################                
#### Main module     
##########################################

# set up results dictionaries
results = {}
for candidate in candidates:
    results[candidate] = {'pass':[],'fail':[],'required':{},'repertoire':{}}

# go through all forms for each juror
for juror in jurors:

    print(juror+' : ')

    # get list of this juror's forms
    filelist = insensitive_glob('*'+juror+'*docx')

    # first deal with adjudication forms
    adjfilelist = [x for x in filelist if x and 'prelim' not in x and 'repertoire' not in x and 'required' not in x]

    # keep track of candidates for which this juror submitted forms
    thiscandlist = []

    for filename in adjfilelist:

        # extract grades from adjudication form
        adjform = Document(filename)

        candidate = get_candnumber(adjform)
        thiscandlist.append(candidate)
        print(candidate,end=' ')
        pieces = get_pieces(adjform)
        reqlist = get_req(adjform)
        grades = get_pf_grades(adjform)
        overall = get_overall(adjform)

        thisdict = copy.deepcopy(results[candidate])

        # store individual piece grades in results
        thisdict = record_grades(thisdict,pieces,grades)
        
        # check for overall pass/fail and consistency
        req_outcomes = [grades[ind]*reqlist[ind] for ind in range(len(pieces))]
        thisdict = record_overall(thisdict,overall,req_outcomes)

        results[candidate] = thisdict

    print('')

    missingcands = [x for x in candidates if x not in thiscandlist]
    if missingcands:
        print('missing form(s) for '+', '.join(missingcands)+' !')
        
    # deal with repertoire form
    repfile = [x for x in filelist if 'repertoire' in x]
    
    if repfile:
        repform = Document(repfile[0])
        # extract values from table in repertoire form and store in dictionary
        results = record_repvotes(repform,results,juror)
    else:
        print("can't find repertoire piece form")

    # deal with required piece form
    reqfile = [x for x in filelist if 'required' in x]

    if reqfile:
        reqform = Document(reqfile[0])
        # extract pieces and choices from required form and store in dictionary
        results = record_reqvotes(reqform,results,juror)
    else:
        print("can't find required piece form")
    

# create summaries of overall scores
make_jurorsummary(results,jurors,voting,conflict)

make_boardsummary(results)

with open('results'+examyear+'.json','w') as fh:
    json.dump(results,fh,indent=4,sort_keys=True)
