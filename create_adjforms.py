'''
Create forms for recording stage adjudications, including
-- adjudication forms for each candidate and juror
-- required piece form for each juror
-- repertoire piece form for each juror
-- pdf program listing summary including all candidates

Created by M. Pan, fall 2020
Edited fall 2021/winter 2022


USAGE:

Move all four .docx templates, the .tsv containing the Google form responses,
and a copy of create_adjforms.py, into the same directory.
Make that directory the current working directory.

$ python3 create_adjforms.py


THINGS TO BE EDITED EACH YEAR:

-- "PARAMETERS" section
-- list of forms_to_make
'''


from __future__ import print_function
from mailmerge import MailMerge
from datetime import datetime
from datetime import date

from docxcompose.composer import Composer
from docx import Document

import glob
import math
import os
import subprocess


#####################
##### UTILITIES #####

def delete_paragraph(paragraph):

    # delete paragraph from docx Document
    
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def delete_row(table,row):
    
    # delete row from table in docx Document

    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def tsv_to_piecedict(tsvfile):

    # convert .tsv of google form responses to dictionary for juror form generation

    with open(tsvfile,'r') as fh:
        tsv = fh.readlines()

    piecedict = {}
        
    if len(tsv) < 2:
        print('no entries in .tsv?')

    else:
        
        entries = [line.strip().split('\t') for line in tsv[1:] if line.strip()]
        candnums = list(set([line[1] for line in entries]))

        for num in candnums:
            thisprog = []
            techflag = 0
            expflag = 0
            programinfo = [line for line in entries if line[1] == num][-1][2:]

            for ind in range(math.floor(len(programinfo)/3)):
                thispiece = {'name':programinfo[3*ind],'comp':programinfo[3*ind+1]}

                if programinfo[3*ind+2] == 'Technical':
                    thispiece['tech'] = True
                    thispiece['exp'] = False
                    techflag += 1
                elif programinfo[3*ind+2] == 'Expressive':
                    thispiece['tech'] = False
                    thispiece['exp'] = True
                    expflag += 1
                else:
                    thispiece['tech'] = False
                    thispiece['exp'] = False

                thisprog.append(thispiece)

            if techflag*expflag != 1:
                print('candidate '+num+' has wrong number of required pieces')
                    
            piecedict[num] = thisprog
                    
                    
    return piecedict
    
        
    
######################
##### PARAMETERS #####    

# names of .docx with mail merge fields
template1 = 'adjform_pf.docx'
template2 = 'overallform.docx'
template3 = 'requiredpieceform.docx'
template4 = 'repertoirepieceform.docx'

# file with responses from program listing form 
tsvfile = '2022 GCNA Carillonneur Exam Recording Program (Responses) - Form Responses 1.tsv' ### !!!CHANGE THIS!!!
piecedict = tsv_to_piecedict(tsvfile)

# year-specific data:
### !!!CHANGE THIS!!!
# year in which exam cycle finishes
examyear = '2022' 

### !!!CHANGE THIS!!!
# all jurors in committee
jurors = ['hunsberger','lee','lehrer','lens','lukyanova','macoska','tam'] 

candidates = sorted(list(piecedict.keys()),key=int)
if '16' in candidates: candidates.remove('16')

### !!!CHANGE THIS!!!
# enforce consistent piece titles in the required piece form
# pick a word likely to appear in any description of this piece,
#    but not in any description of any other piece
# associate it with the specific title to be used in the required piece form
req_piece_std_format = {
    'burlesca': 'Burlesca',
    'chanticleer': 'Call of the Chanticleer',
    'fugue': 'Cortege and Fugue',
    'arlington': 'Poème pour Arlington',
    'valse': 'Valse Romantique',
    'braes': "Ye Banks and Braes"}


#############################
##### FORMS TO GENERATE #####

# 'adj' : adjudication forms
# 'rep' : repertoire piece forms
# 'req' : required piece forms
# 'prog' : program listings

forms_to_make = ['rep','req','adj','prog']
    

##############################
##### ADJUDICATION FORMS #####

if 'req' in forms_to_make:
    reqpiecedict = {'exam_year':examyear}

for candidate in candidates:

    piecelist = piecedict[candidate]
    
    if 'req' in forms_to_make:
        
        # setup dictionary of all candidates' required pieces for use in required piece form    
        try:
            tech = [x['name'] for x in piecelist if x['tech'] == True][0]
            exp = [x['name'] for x in piecelist if x['exp'] == True][0]
            techkey = [x for x in req_piece_std_format.keys() if x in tech.lower()]
            expkey = [x for x in req_piece_std_format.keys() if x in exp.lower()]
            if techkey:
                tech = req_piece_std_format[techkey[0]]
            if expkey:
                exp = req_piece_std_format[expkey[0]]
        except:
            print('missing required piece for candidate '+candidate)
            tech = ''
            exp = ''
        reqpiecedict['tech'+candidate] = tech
        reqpiecedict['exp'+candidate] = exp

        
    if 'adj' in forms_to_make:
    
        # create adjudication forms (1 form per piece) for this candidate  
        for ind in range(len(piecelist)):

            piece = piecelist[ind]
            piecefields = {
                'candidate_number': candidate,
                'piece_name': piece['name'],
                'composer_name': piece['comp'],
            }
            if piece['tech'] or piece['exp']:
                piecefields['req'] = 'Yes'
            else:
                piecefields['req'] = 'No'

            doc = MailMerge(template1)
            doc.merge(**piecefields)
            doc.write(examyear+'_'+candidate+str(ind+1)+'.docx')

        # create overall pass/fail page
        overall = {'candidate_number': candidate}
    
        enddoc = MailMerge(template2)
        enddoc.merge(**overall)
        enddoc.write(examyear+'_'+candidate+'end.docx')

        # combine all adjudication forms and overall pass/fail page
        master = Document(examyear+'_'+candidate+'1.docx')
        composer = Composer(master)

        for ind in range(len(piecelist)-1):
            section = Document(examyear+'_'+candidate+str(ind+2)+'.docx')
            composer.append(section)
        
        endsection = Document(examyear+'_'+candidate+'end.docx')
        composer.append(endsection)

        # save one copy for each juror
        for juror in jurors:

            thisadjfile = examyear+'_candidate'+candidate+'_'+juror+'.docx'
            composer.save(thisadjfile)

            if not os.path.exists(juror):
                os.mkdir(juror)

            os.replace(thisadjfile,os.path.join(juror,thisadjfile))

        for adjfile in glob.glob(examyear+'_'+candidate+'*.docx'):
            os.remove(adjfile)

        
###############################
##### REQUIRED PIECE FORM #####

if 'req' in forms_to_make:

    # create required piece form        
    reqpiecedoc = MailMerge(template3)
    reqpiecedoc.merge(**reqpiecedict)
    reqpiecedoc.write(examyear+'_reqpieceform.docx')

    # delete fields for nonexistent candidates
    reqpieceform = Document(examyear+'_reqpieceform.docx')
    delcounter = 0
    for paragraph in reqpieceform.paragraphs:
        if delcounter > 0:
            delete_paragraph(paragraph)
            delcounter = delcounter-1
        else:
            candnum = paragraph.text.split('\t')[0]
            if 'Candidate' in candnum and candnum[-2:].strip().isnumeric() and candnum[-2:].strip() not in candidates:
                delete_paragraph(paragraph)
                delcounter = 2

    # save one copy per juror
    for juror in jurors:

        reqfile = examyear+'_requiredpieceform_'+juror+'.docx'
        reqpieceform.save(reqfile)

        if not os.path.exists(juror):
            os.mkdir(juror)

        os.replace(reqfile,os.path.join(juror,reqfile))

    os.remove(examyear+'_reqpieceform.docx')
        

#################################
##### REPERTOIRE PIECE FORM #####

if 'rep' in forms_to_make:

    # setup dictionary of required pieces
    repertoiredict = {'exam_year':examyear}
    for ind in range(len(candidates)):
        repertoiredict['candnum'+str(ind+1)] = candidates[ind]

    # create repertoire piece form
    repertoiredoc = MailMerge(template4)
    repertoiredoc.merge(**repertoiredict)
    repertoiredoc.write(examyear+'_reppieceform.docx')

    # delete fields for nonexistent candidates
    repertoireform = Document(examyear+'_reppieceform.docx')
    piecetable = repertoireform.tables[0]
    for row in piecetable.rows:
        if not row.cells[0].text:
            delete_row(piecetable,row)

    # save one copy per juror
    for juror in jurors:

        repfile = examyear+'_repertoirepieceform_'+juror+'.docx'
        repertoireform.save(repfile)

        if not os.path.exists(juror):
            os.mkdir(juror)

        os.replace(repfile,os.path.join(juror,repfile))
        
    os.remove(examyear+'_reppieceform.docx')
    

############################
##### PROGRAMS LISTING #####

if 'prog' in forms_to_make:

    # create LaTeX file with candidate programs; start with LaTeX front matter
    preamble = [
        '\\documentclass[10pt]{article}\n', \
        '\n', \
        '\\usepackage{parskip,array}\n', \
        '\\usepackage[scaled=.9]{helvet}\n', \
        '\\usepackage[T1]{fontenc}\n', \
        '\n', \
        '\\addtolength{\\topmargin}{-.9in}\n', \
        '\\addtolength{\\oddsidemargin}{0in}\n', \
        '\\addtolength{\\oddsidemargin}{-1in}\n', \
        '\\addtolength{\\textwidth}{2in}\n', \
        '\\addtolength{\\textheight}{1.7in}\n', \
        '\n', \
        '\\renewcommand\\familydefault{\\sfdefault}\n', \
        '\\renewcommand{\\arraystretch}{1.1}\n', \
        '\n', \
        '\\begin{document}\n', \
        '\n', \
        '\\begin{LARGE}\n', \
        '\\noindent {\\bf '+examyear+' Recording Program Listings}\\bigskip\\hfill\n', \
        '\\end{LARGE}\n', \
        '\n', \
        ]

    proglines = preamble

    # index for individual program listings
    pieceindex = 'abcdefgh'
    
    for candidate in candidates:

        # create tabular env for each candidate's program
        thisprog = piecedict[candidate]
        thisproglist = ['\\begin{tabular}{p{0.13\\textwidth}p{0.02\\textwidth}<{\\raggedleft\\arraybackslash}p{0.38\\textwidth}<{\\raggedright\\arraybackslash}p{0.25\\textwidth}<{\\raggedright\\arraybackslash}p{0.03\\textwidth}}\n']
        piececount = 0

        for piece in thisprog:
            piecestr = '& ' + pieceindex[piececount]+')& ' + piece['name']+'& ' + piece['comp']+'& '
            if piececount == 0:
                piecestr = '{\\bf Candidate '+candidate+'}' + piecestr
                
            if piece['tech']:
                piecestr += '(T)\\\\\n'
            elif piece['exp']:
                piecestr += '(E)\\\\\n'
            else:
                piecestr += '\\\\\n'
                
            thisproglist.append(piecestr)
            piececount += 1

        thisproglist.append('\\end{tabular}\\medskip\n\n')
        
        proglines += thisproglist

    proglines += ['\\end{document}\n']

    # save .tex and run pdflatex
    progfile = examyear+'_candidate_programs.tex'
    with open(progfile,'w') as fh:
        for line in proglines:
            _ = fh.write(line)

    subprocess.run('pdflatex '+progfile+' > pdflatex.log 2>&1',shell=True)
    
